#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import print_function
import os
import sys
import json
import glob
import argparse

import requests
from docx import Document
import bs4
from bs4 import BeautifulSoup

"""
Parses one or more JSON files and outputs well-formatted DOC(X) file(s).

Types of nodes:
c(hunk): chunk of text; may be the entire text, a sentence/unit, clause, phrase, etc...
l(emma): text token (inflected, conjugated, etc.) along with its lemmatization ("dictionary" form and specific meaning)
d(iscontinuity): line break, surface transition, or damage
"""

closing_punct_mirror = {
    "]": "[",
    ")": "(",
    "›": "‹",
    "»": "«",
    ")›": "‹(",
}

# Used to transform eg. a2 to á in the final result docx.
vowel_subscript_to_accent_map = {
    "a": {
        "two": "á",
        "three": "à",
    },
    "e": {
        "two": "é",
        "three": "è",
    },
    "i": {
        "two": "í",
        "three": "ì",
    },
    "u": {
        "two": "ú",
        "three": "ù",
    },
}

VERBOSE_FLAG = False


def print_if_verbose(msg):
    global VERBOSE_FLAG
    if VERBOSE_FLAG:
        print(msg)


class JsonLoader(object):
    """
    Class to read from a filename/pathname containing one or more JSON files
    and make accessible a list of dicts (one per JSON/exemplar).
    """
    def __init__(self, original_path):
        print_if_verbose("Using encoding {0}".format(sys.stdout.encoding)) # cp1252; can't process some UTF-8 stuff because windoze :(

        self.catalogue_dict = None
        self.json_paths = self._get_file_paths(original_path)
        self.json_dicts = self._load_json_dicts()
        self.q_number = os.path.basename(original_path).split(".json")[0]

    def _get_file_paths(self, original_path):
        """Returns a list of one or more ORACC JSON files to read. Passing
        in a directory will result in a list of all JSONs in that directory.
        Args:
            original_path (str): Relative or absolute path to either one
                JSON file or a directory of JSON files
        Returns:
            list: List of one or more path strings to JSON files
        Raises:
            Exception: if original_path is invalid
        """
        if os.path.isfile(original_path):
            return [original_path]
        elif os.path.isdir(original_path):  # glob the files directly in dir
            path = os.path.join(original_path, "*.json")
            return glob.glob(path)
        else:
            raise Exception("Invalid path specified! "
                            "Please ensure that your given path points to either "
                            "a .json or a directory containing .json files.")

    def _load_json_dicts(self):
        """Loads one or more ORACC JSON files into one or more python dicts.
        If a JSON file is unable to be read, its corresponding dict will be empty,
        save for its original path.
        Returns:
            list (dict): loaded from the raw JSON files
        """
        json_dicts = []
        for json_path in self.json_paths:
            try:
                with open(json_path, encoding="utf_8_sig") as fd:
                    raw_str = fd.read()
                    json_dict = json.loads(raw_str)
                    q_number = json_dict["textid"] # aka. CDLI number

                    if not self.catalogue_dict:
                        self.catalogue_dict = self._get_catalogue_json(json_path)

                    # Add in additional data to JSONs, mostly from their catalog
                    {'collection': 'Iraq Museum, Baghdad, Iraq',
                     'designation': 'Unidentified Suhu 1007',
                     'display_name': 'Suhu Unidentified Suhu 1006',
                     'museum_no': 'IM 096751',
                     'popular_name': 'RIMB 2 S.0.0.1006',
                     'primary_publication': 'Unidentified Suhu 1006'}

                    json_dict["original_path"] = json_path
                    q_catalogue = self.catalogue_dict["members"][q_number]

                    json_dict["museum_no"] = q_catalogue.get("museum_no") # seen in SAAO, SUHU
                    if json_dict["museum_no"] == "IM -": # duds, seen in SAAO
                        json_dict["museum_no"] = ''

                    json_dict["exemplars"] = q_catalogue.get("exemplars") # Seen in RINAP, RIBO
                    json_dict["collection"] = q_catalogue.get("collection") # same as above; add as supplemental info

                    json_dict["primary_publication"] = q_catalogue["primary_publication"] # eg. Esarhaddon 088, Tiglath-pileser III 01, SAA 19 215,


                    if json_dict["museum_no"]:
                        json_dict["ochre_title"] = json_dict["museum_no"]
                    else:
                        json_dict["ochre_title"] = "(PUB) " + json_dict["primary_publication"]
                    # TODO NOTE idea: have text file with real museum info for RINAP/RIBO lined up with the q-nums. I don't know which is the real publication info anymore
                    # even just a Q-num textfile + hotkey to prepopulate name of doc can help...

                    json_dict["docx_name"] = q_number

                    json_dicts.append(json_dict)
            except Exception as e:
                print("Could not load {0} to dict: {1}".format(json_path, e))
                print("If this is an encoding error, check that the venv is based on py3, not py2")
                json_dicts.append({
                    "original_path": json_path,
                })
        return json_dicts

    def get_json_dicts(self):
        return self.json_dicts # TODO make this into property

    def _get_catalogue_json(self, json_path):
        """Gets output of reading from json_path/../catalogue.json.
        """
        try:
            catalogue_path = os.path.join(os.path.dirname(json_path), "..", "catalogue.json")
            return self._read_json_dict(catalogue_path)
        except Exception as e:
            print("Unable to find catalogue.json at {0}.".format(catalogue_path))
            raise e

    def _read_json_dict(self, filename):
        with open(filename) as fd:
            raw_str = fd.read()
            return json.loads(raw_str)


class JsonParser(object):
    """
    Class to take in a local JSON file and output a docx.
    """
    def __init__(self, json_dict, output_directory):
        self.cdl_dict = json_dict
        self.output_directory = output_directory
        self.l_reflist = []  # for repeat nodes
        self.traversed_first_c_sentence_node = False  # TODO don't need this anymore if we're not using c-labels
        self.found_obverse_or_reverse_d_node = False  # trigger on first type=discourse c-node?

        self.q_number = json_dict.get("textid", None)  # NOTE: this can be a P-number, q-number is misleading...
        self.project = json_dict.get("project", None)

        self.exemplars = json_dict.get("exemplars") # extra from JSONReader
        self.primary_publication = json_dict.get("primary_publication") # eg. Esarhaddon 088
        self.museums = json_dict.get("collection") # eg. British Museum, London, UK

        self.soup = None
        self.has_aramaic = False

    def run(self):
        """Loads and parses given ORACC JSON, then saves the pieced-together
        text into a docx on disk.
        """
        if not self.q_number:
            print_if_verbose("Skipping malformed JSON from {0}:".format(self.cdl_dict["original_path"]))
            print_if_verbose(self.cdl_dict)
            return
        print_if_verbose(
            "Parsing textid {0} from project {1}".format(self.q_number, self.cdl_dict["project"])
        )
        doc = Document()
        res = self.parse_json(doc)
        self.print_doc(doc)
        self.save_docx(doc)

    def parse_json(self, doc):
        """Walks through the JSON object and pieces together all the lemmas.
        No new sections like obverse/reverse yet. TODO complete docstring

        dict[cdl] is list
        dict[cdl][3] (or whatever index) has node == c; start parsing there
        dict[cdl][3][cdl] contains list of dicts with actual lemmas/line breaks
        """
        if not self.cdl_dict:
            return

        if self.cdl_dict["type"] != "cdl":
            print("Not a CDL-type JSON!\n")
            return

        nodes = self.cdl_dict["cdl"]
        for c_node in nodes:
            chunk = c_node
            res = self.traverse_c_node(chunk, doc)

    def save_docx(self, doc):
        """Attempt to save the resulting docx file to current directory where
        script originally ran. Resulting docx will either be named after its
        Q-number textid or its exemplar sources.
        Args:
            #textid (str): ID of original JSON dict; basis of save name
                (eg. Q003456 -> Q003456.docx)
            doc (docx.Document): fully assembled docx object to be saved
        """
        try:
            # Check first to make sure there's anything worth saving, eg. an empty JSON
            p_text = ""
            for p in doc.paragraphs:
                for r in p.runs:
                    p_text += r.text
            p_text = p_text.strip()
            if not p_text or p_text == "Obverse" or p_text == "Text":
                print_if_verbose("No text in this docx- skipping save!")
                return

            # Otherwise, go on and save it
            docx_name = self._get_docx_name_to_save(self.cdl_dict["docx_name"]) + ".docx"
            if self.has_aramaic:
                docx_name = "(arc) " + docx_name
            docx_path = os.path.join(self.output_directory, docx_name)
            doc.save(docx_path)
            print("Saved docx in {0}".format(docx_path))
        except Exception as e:
            print("Couldn't save docx! {0}".format(e))

    def _get_docx_name_to_save(self, name):
        """Check if docx_name would be unique before saving. If not, append the appropriate
        number identifier. eg. "My Exemplar Name" -> "My Exemplar Name (1)", assuming
        "My Exemplar Name" exists.
        """
        number_id = 0
        original_name = name
        docx_name = name

        while True:
            if not os.path.exists(docx_name + ".docx"):
                return docx_name
            print("{0} exists, continuing".format(docx_name))
            number_id += 1
            docx_name = original_name + " ({0})".format(number_id)

    def print_doc(self, doc):
        """Utility function to print resulting fully assembled text to console.
        There will be no formatting such as italics. Super/subscripts may depend
        based on your terminal of choice.

        Args:
            doc (docx.Document): fully assembled text result to be printed
        """
        full_left_brackets = 0
        full_right_brackets = 0
        partial_left_brackets = 0
        partial_right_brackets = 0
        for p in doc.paragraphs:
            s = ""
            for r in p.runs:
                s += r.text
                if "[" in r.text:
                    full_left_brackets += 1
                elif "]" in r.text:
                    full_right_brackets += 1
                elif "⸢" in r.text:
                    partial_left_brackets += 1
                elif "⸣" in r.text:
                    partial_right_brackets += 1
            print_if_verbose(s)

        # Seeing if we're balanced or not
        #print_if_verbose("[ count: {0}".format(full_left_brackets))
        #print_if_verbose("] count: {0}".format(full_right_brackets))
        #print_if_verbose("⸢ count: {0}".format(partial_left_brackets))
        #print_if_verbose("⸣ count: {0}".format(partial_left_brackets))
        #assert(full_left_brackets == full_right_brackets) # TODO restore if you remove exemplar name
        #assert(partial_left_brackets == partial_right_brackets) # TODO ditto
        print_if_verbose("--------------------------------------\n")

    def traverse_c_node(self, c_dict, doc, first_c_node=False):
        """Decides what to do for each of the nodes in a c-node's node list.
        Will further traverse down a C(hunk), D(iscontinuity), or L(emma)
        node as needed.

        Args:
            c_dict (dict): Corresponds to the very first C node encountered
                in an ORACC JSON. Contains more C, D, or L nodes that may
                further be nested.
            doc (docx.Document): docx object to append lemmas to
        """
        if not c_dict.get("id", ""):
            print_if_verbose("No id for this c-node- returning!")
            return

        print_if_verbose("At c-node {0}".format(c_dict["id"]))

        # If we've reached the point of starting the actual text, ie. no D-nodes
        # with "obverse" or "reverse" have been encountered so far, we still need
        # some header, so we put in "Text"
        if c_dict["type"] == "discourse":
            if not self.found_obverse_or_reverse_d_node:
                p = doc.add_paragraph()
                p.add_run("Text")
                doc.add_paragraph()
                self.found_obverse_or_reverse_d_node = True

        for node in c_dict["cdl"]:
            if node["node"] == "c":
                self.traverse_c_node(node, doc, first_c_node=False)
            elif node["node"] == "d":
                self.parse_d_node(node, doc)
            elif node["node"] == "l":
                self.parse_l_node(node, doc)
            else:
                print_if_verbose("Unknown node type for node {0}".format(node))

    def parse_d_node(self, d_dict, doc):
        """Parses a D(iscontinuity) node and adds paragraphs to doc as needed.
        Types of d-node values:
          - line-start
          - obverse
          - reverse
          - excised
          - object
          - surface
          - tablet
          - punct
          - nonx/nonw
        """
        assert(d_dict["node"] == "d")
        d_type = d_dict["type"]

        if d_type == "line-start":
            # If there was eg. a lacuna before, don't insert a newline.
            # We want to keep a text flowing so that there's no empty lines
            # aside from before things like "Reverse"
            try:
                # Assemble any text from the last run of last paragraph
                p_text = ""
                for r in doc.paragraphs[-1].runs:
                    p_text += r.text

                # If text contains prior D-node, make sure to not add a space!
                if "Text" in p_text or "Obverse" in p_text or "Reverse" in p_text or "column" in p_text:
                    print_if_verbose("Last line was a D-node header '{0}', skipping line-start".format(p_text))

                elif p_text.strip() == "":  # a whitespace-only paragraph
                    print_if_verbose("Last paragraph was empty, skipping line-start")

                else:
                    print_if_verbose("Last paragraph was NOT empty. Applying line-start newline.")
                    doc.add_paragraph()
            except Exception as e:
                print("Couldn't get last run before this line-start: {0}".format(e))
            # NOTE: disabled below since this was adding eg. custom line numbers that OCHRE won't be able to parse
            # Seems like custom header labels should come from C-nodes...? If at all?
            #p.add_run(d_dict.get("label", ""))  # eg. Inscription_A 1 in rinap4/Q003347

        elif d_type == "obverse":
            p = doc.add_paragraph()
            # there'll be at least 2 paragraphs already if "Text" present
            # if "obverse" in the middle and not at start, needs extra newline
            if len(doc.paragraphs) >= 2:
                p = doc.add_paragraph()
            p.add_run("Obverse")
            doc.add_paragraph()
            self.found_obverse_or_reverse_d_node = True

        elif d_type == "reverse":
            # Needs extra newline before it, unlike obverse, since it comes later on in texts
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Reverse")
            doc.add_paragraph()
            self.found_obverse_or_reverse_d_node = True

        elif d_type == "punct":
            p = doc.paragraphs[-1]
            p.add_run(d_dict["frag"])
            p.add_run(d_dict.get("delim", ""))

        # TODO: below is WIP
        # what happens if the very first element of the paragraph is excised?
        elif d_type == "excised" and "frag" in d_dict:
            assert(len(doc.paragraphs) > 0)
            self._add_excised_d_node(d_dict, doc.paragraphs[-1])

        elif d_type == "excised" and "frag" not in d_dict:
            print_if_verbose("Excised node without frag:")
            print_if_verbose(d_dict)

        elif d_type == "nonx" or d_type == "nonw" or d_type == "object" or d_type == "surface":
            pass

        else:
            print_if_verbose("Unknown or noop d-value {0}".format(d_type))

    def _add_excised_d_node(self, d_dict, paragraph):
        """Add a D-node of type "excised". These nodes don't come with the same members/metadata as L-nodes,
        even if the node's contents have more than one sign in it. Doesn't replace subscript #'s or convert 2/3
        subscripts to accented marks- apparently OCHRE understands that fine.

        TODO: see how well this works with SAAO/SAA16 or CAMS that may have eg. [<<BU>>], <<BU>>], or {<<d}60.
        TODO: test with eg. "{<<uru}arba-il₃>>" like in P334914 in saao
        TODO reuse this with SAAO fragments that aren't d-nodes?
        """
        frag = d_dict["frag"].replace("<<", "«").replace("<", "‹").replace(">>", "»").replace(">", "›").replace("$", "")

        det_start_index = frag.find("{")
        det_end_index = frag.find("}")

        # if only } is detected, let's just start off assuming it's a determinative until we meet }
        if det_start_index == -1 and det_end_index > -1:
            det_mode = True
        else:
            det_mode = False

        for char in frag:
            if char == "{":  # begin det mode, ignore char
                det_mode = True
            elif char == "}":  # end det mode, ignore char
                det_mode = False
            elif char.isalpha() or char.isdigit(): # Sumerian or Akkadian, or a subscript #
                if det_mode and char.islower() and char != "m" and char != "d": # NOTE this assumes only Sumerian determinatives
                    char = char.capitalize() # tested and should work fine with eg. Ğ.
                r = paragraph.add_run(char)
                if char.islower():  # Akkadian - set italics
                    r.italic = True
                if det_mode:
                    r.superscript = True
            else:  # symbol, probably like - or [ or ], or << <
                paragraph.add_run(char)

        paragraph.add_run(d_dict.get("delim"))
        print_if_verbose("Added excised D-node {0}".format(d_dict['frag']))
        print_if_verbose(d_dict)

    def parse_l_node(self, l_dict, doc):
        """Gets L(emma) node text, formats it, and adds it to the last paragraph of doc.
        L nodes may contain either an Akkadian or Aramaic lemma.
        Example of Akkadian L node which this function will work on (in rinap/rinap1/corpusjson/Q003414.json):
        {
            "node": "l",
            "frag": "{d}[(...)",
            "id": "Q003414.l000a0",
            "ref": "Q003414.2.5",
            "inst": "u",
            "f": {
                "lang": "akk",
                "form": "{d}x",
                "delim": " ",
                "gdl": [
                    {
                        "det": "semantic",
                        "pos": "pre",
                        "seq": [
                          {
                            "v": "d",
                            "gdl_utf8": "𒀭",
                            "id": "Q003414.2.5.0"
                          }
                        ]
                    },
                    {
                        "x": "ellipsis",
                        "id": "Q003414.2.5.1",
                        "breakStart": "1",
                        "statusStart": "Q003414.2.5.1",
                        "break": "missing",
                        "o": ")"
                    }
                ],
                "pos": "u"
            }
        }
        Args:
            l_dict (dict): dict version of an L node above
            doc (docx.Document): document object to append lemma to
        """
        print_if_verbose("At L-node {0}".format(l_dict["ref"]))
        # Check if this frag's already been added
        ref = l_dict["ref"]
        if ref in self.l_reflist:
            print_if_verbose("Already added ref {0}, skipping".format(ref))
            return
        else:
            self.l_reflist.append(ref)

        last_paragraph = doc.paragraphs[-1]
        lang = l_dict.get("f").get("lang")

        if lang == "arc":  # eg. Aramaic
            print_if_verbose("Adding Aramaic fragment")
            self._add_aramaic_frag(l_dict, last_paragraph)
            return
        elif lang == "qcu-949":  # seemingly English...
            print_if_verbose("Not adding English fragment...")
            return
        elif lang != "akk" and lang != "akk-949" and \
                lang != "akk-x-neoass" and lang != "sux" and \
                lang != "akk-x-neobab":
            print_if_verbose("Unrecognized language {0}".format(lang))

        # TODO: I dunno what to do for this...
        gdl_list = l_dict.get("f", "").get("gdl", "")
        if not gdl_list:
            # This seems to happen only for SAAO/Suhu: no gdl in the f member of L-node
            # However, contents of f is not usable- it's often an assembled Akkadian word
            # rather than transliterated version (eg. bilticu vs. GUN-cu).
            # We'll have to use the online version at this point using the ref #
            print("INCOMPLETE TEXT starting at {0}- scraping web equivalent at http://oracc.museum.upenn.edu/{1}/{2}".format(l_dict["ref"], self.project, self.q_number))
            print_if_verbose("Raw fragment is {0}".format(l_dict["frag"]))
            self._scrape_incomplete_l_node(l_dict["ref"], last_paragraph)
            return

        for index in range(len(gdl_list)):
            node_dict = gdl_list[index]
            if "s" in node_dict:
                self._add_logogram(node_dict, last_paragraph)
            elif "v" in node_dict:
                self._add_continuing_sign_form(node_dict, last_paragraph)
            elif "det" in node_dict:
                # NOTE: if there's 2 determinatives stuck next to each other,
                # need to separate them with space or something else
                # since OCHRE will otherwise attempt to look up
                # eg. "md" instead of "m" and "d" dets separately
                try:
                    if len(gdl_list) > index + 1 and "det" in gdl_list[index + 1]:
                        self._add_determinative(node_dict, last_paragraph, add_dot_delim=True)
                        print_if_verbose("Added first in set of multiple DETs")
                    else:
                        self._add_determinative(node_dict, last_paragraph)
                except Exception as e:
                    print("Looks like a single determinative, not 2 stuck together! Exception: {0}".format(e))
                    raise e # NOTE keeping this around for debug purposes; this ideally should never hit
            elif "gg" in node_dict:
                self._add_logogram_cluster(node_dict, last_paragraph)
            elif "x" in node_dict:
                self._add_ellipsis(node_dict, last_paragraph)
            elif "n" in node_dict:
                self._add_number(node_dict, last_paragraph)
            elif "q" in node_dict:
                # TODO add dedicated function
                d_dict = {
                    "frag": node_dict["q"].replace("|", ""),
                    "delim": node_dict.get("delim"),
                }
                self._add_excised_d_node(d_dict, last_paragraph)
                print_if_verbose("Added qualified element {0} via parse_l_node".format(node_dict["q"]))
            elif "c" in node_dict:
                # TODO add dedicated function
                c_frag = node_dict["c"].replace("|", "")
                last_paragraph.add_run(c_frag + node_dict.get("delim"))
                print_if_verbose("Added composite fragment {0}".format(node_dict['c']))
            elif "mods" in node_dict:
                self._add_pre_frag_symbols(node_dict, last_paragraph)
                frag = node_dict["form"]
                r = last_paragraph.add_run(frag)
                if frag.islower():
                    r.italic = True
                self._add_post_frag_symbols(node_dict, last_paragraph)
                last_paragraph.add_run(node_dict.get("delim", ""))
                print_if_verbose("Added mods L-node {0}".format(node_dict["form"]))
            else:
                print_if_verbose("Unknown l-node {0}".format(node_dict))
        last_paragraph.add_run(l_dict["f"].get("delim", "")) # TODO still needed?

    def _scrape_incomplete_l_node(self, ref_id, paragraph):
        """For L-nodes that have no gdl_dict and must rely on their online counterparts in ORACC
        to get inputted properly. Usually from ribo/babylon6.
        eg. from http://oracc.museum.upenn.edu/rinap/rinap1/Q003418/html:
        <span class="w N " id="Q003418.5.12">
          <a class="cbd " >
            ⸢
            <span class="sign sux ">UD</span>
            ⸣.
            <span class="sign sux ">MEŠ</span>
          </a>
        </span>
        NOTE: ref_id isn't guaranteed to be in the web equivalent, let's ignore it if it's missing
        """
        url = "http://oracc.museum.upenn.edu/{0}/{1}".format(self.project, self.q_number)

        if not self.soup: # lazy load
            self.soup = BeautifulSoup(requests.get(url).content, "html.parser")

        parent = self.soup.find("span", {"id": ref_id}) # Assume only 1 span with this ID

        if not parent:
            print_if_verbose("Skipping this ID scrape - id {0} not in web equivalent".format(ref_id))
            return

        if parent.a: # eg. http://oracc.museum.upenn.edu/suhu/Q006238 has no <a> below
            parent = parent.a

        # TODO: replace <<, <, etc. first. Is it needed at all though? d-nodes are what have << >> usually

        for snippet in parent:
            if type(snippet) is bs4.element.Tag: # is a <span> or <sup>
                if snippet.text == "?": # Online ORACC has superscript ?, but we want non-superscript (?)
                    r = paragraph.add_run("(?)")
                    print_if_verbose("Adding ? sup snippet as non-superscript (?)")
                    continue
                r = paragraph.add_run(snippet.text)
                if snippet.name == "sup":
                    r.font.superscript = True
                    print_if_verbose("Adding scraped determinative {}".format(snippet.text))
                elif snippet.name == "span" and "akk" in snippet["class"]:
                    r.italic = True
                    print_if_verbose("Adding scraped Akkadian {}".format(snippet.text))
                else:
                    print_if_verbose("Adding scraped Sumerian {}".format(snippet.text))
            elif type(snippet) is bs4.element.NavigableString: # is just filler chars like [
                r = paragraph.add_run(snippet)
                print_if_verbose("Adding scraped etc character {}".format(snippet))
        paragraph.add_run(" ") # assumed delim afterwards

    def _add_aramaic_frag(self, l_node, paragraph):
        """Adds Aramaic fragment to current paragraph with all needed formatting.
        Example of Aramaic L node which this function will work on (in rinap/rinap1/corpusjson/Q003633.json):
        {
            "node": "l",
            "frag": "mnn",
            "id": "Q003633.l00075",
            "ref": "Q003633.1.1",
            "inst": "%arc:mnn=",
            "f": {
                "lang": "arc",
                "form": "mnn"
            }
        }
        Args:
            l_node (dict): Aramaic lang node to be added
            paragraph (docx.text.paragraph.Paragraph): paragraph to add Aramaic fragment to
        """
        frag = l_node.get("frag", "")
        for char in frag:
            r = paragraph.add_run(char)
            if char.isalpha():
                r.italic = True
        # Aramaic nodes have no "delim", but should be separated with space
        paragraph.add_run(" ")
        self.has_aramaic = True

    def _add_continuing_sign_form(self, gdl_node, paragraph):
        """Adds eg. tu- to the current paragraph.
        eg.
        {
          "v": "a",
          "queried": "1",
          "ho": "1",
          "delim": "-"
        }
        """
        self._add_pre_frag_symbols(gdl_node, paragraph)

        # Actual sign/word fragment
        word = self._convert_2_or_3_subscript(gdl_node["v"])
        r = paragraph.add_run(word)
        if word.islower():
            r.italic = True
        ##print_if_verbose("Added continuing sign {0}".format(word))

        self._add_post_frag_symbols(gdl_node, paragraph)

    def _add_determinative(self, gdl_node, paragraph, add_dot_delim=False):
        """Adds determinative to given paragraph and adds necessary styling.
        Example of a node that this would work on (from rinap/rinap1/corpusjson/Q003414.json):
        {
            "det": "semantic",
            "pos": "post",
            "seq": [
                {
                    "s": "KI",
                    "gdl_utf8": "𒆠",
                    "id": "Q003414.2.2.2",
                    "role": "logo",
                    "logolang": "sux"
                }
            ]
        }

        Args:
            gdl_node (dict): dict representation of a "det" node. These will
                always occur as a member of a L(emma) node's gdl list.
            paragraph (docx.text.paragraph.Paragraph): Paragraph object
                that will append the new determinative
        """
        assert(gdl_node.get("det", "") == "semantic" or
               gdl_node.get("det", "") == "phonetic")

        if gdl_node["pos"] == "pre" or gdl_node["pos"] == "post":
            det_node = gdl_node["seq"][0] # TODO assumes seq dict only has 1 member- seems true so far...

            self._add_pre_frag_symbols(det_node, paragraph) # TODO document why det node passed in

            # Add the determinative to paragraph with needed stylings
            if "s" in det_node: # traditional DET node
                det = det_node["s"]
                print_if_verbose("Adding regular det form {0}".format(det))
            elif "v" in det_node: # alt to first; esp. for ones like m or d
                det = det_node["v"]
                print_if_verbose("Adding alternative det form {0}".format(det))
            elif "mods" in det_node: # eg. LU2~v
                det = det_node["form"]
                print_if_verbose("Adding MODS det form {0}".format(det))
            elif "n" in det_node: # numeral det, eg. 1(dic)
                # The 1 det is actually the same as m... compare saa19/x900013 online and in json
                det = det_node["form"]
                if det == "1":
                    det = "m"
                print_if_verbose("Adding numeral det {0}".format(det))
            else:
                print_if_verbose("Unknown DET type: {0}".format(det_node))

            det = self._convert_2_or_3_subscript(det)
            r = paragraph.add_run(det)
            r.font.superscript = True

            self._add_post_frag_symbols(det_node, paragraph)
            if add_dot_delim:  # if there's another det right after this
                r = paragraph.add_run(".")  # TODO use . or space? Space looked a bit weird, so let's try .
                r.font.superscript = True
                print_if_verbose("Added extra . delim for double determinative:")
                print_if_verbose(gdl_node)
        else:
            print_if_verbose("Unknown determinative position {0}".format(gdl_node["pos"]))

    def _add_logogram(self, gdl_node, paragraph):
        """Adds a standalone logogram to current paragraph, eg. LUGAL.
        Example L(emma) node (from rinap/rinap1/corpusjson/Q003627.json):
        {
            "node": "l",
            "frag": "LUGAL",
            "id": "Q003627.l00054",
            "ref": "Q003627.2.4",
            "inst": "šarri[king]N +.",
            "sig": "@rinap/rinap1%akk:LUGAL=šarru[king//king]N'N$šarri",
            "f": {
              "lang": "akk",
              "form": "LUGAL",
              "gdl": [ # begin gdl_node
                {
                  "s": "LUGAL",
                  "gdl_utf8": "𒈗",
                  "id": "Q003627.2.4.0",
                  "role": "logo",
                  "logolang": "sux"
                }
              ], # end gdl_node
              ...
        }
        """
        assert(gdl_node.get("s", ""))
        if gdl_node.get("role", "") != "logo":
            print_if_verbose("Non-logo logogram found! {0}".format(gdl_node["s"]))

        self._add_pre_frag_symbols(gdl_node, paragraph)

        # Add actual logogram
        logogram = self._convert_2_or_3_subscript(gdl_node["s"])
        paragraph.add_run(logogram)
        ##print_if_verbose("Added logogram {0}".format(logogram))

        self._add_post_frag_symbols(gdl_node, paragraph)

    def _add_logogram_cluster(self, gdl_node, paragraph):
        """Adds >1 logograms to current paragraph, eg. GIC.TUG.PI.
        Example L(emma) node containing gdl node "gdl" (from rinap/rinap1/corpusjson/Q003627.json):
        {
            "node": "l",
            "frag": "MA.NA",
            "id": "Q003627.l00052",
            "ref": "Q003627.2.2",
            "inst": "manê[a unit of weight]N",
            "sig": "@rinap/rinap1%akk:MA.NA=manû[unit//a unit of weight]N'N$manê",
            "f": {
              "lang": "akk",
              "form": "MA.NA",
              "delim": " ",
              "gdl": [ # begin gdl_node
                {
                   "gg": "logo",
                   "gdl_type": "logo",
                   "group": [
                     {
                       "s": "MA",
                       "gdl_utf8": "𒈠",
                       "id": "Q003627.2.2.0",
                       "role": "logo",
                       "logolang": "sux",
                       "delim": "."
                     },
                     {
                       "s": "NA",
                       "gdl_utf8": "𒈾",
                       "id": "Q003627.2.2.1",
                       "role": "logo",
                       "logolang": "sux"
                     }
                  ]
                }
            ], # end gdl_node
            ...
        }
        """
        logo_group_dict = gdl_node["group"]
        for logo_dict in logo_group_dict:
            if "s" in logo_dict:
                self._add_logogram(logo_dict, paragraph)
            elif "det" in logo_dict:
                self._add_determinative(logo_dict, paragraph)
            elif "v" in logo_dict:
                self._add_continuing_sign_form(logo_dict, paragraph)
            elif "n" in logo_dict:
                self._add_number(logo_dict, paragraph)
            elif "gg" in logo_dict:
                self._add_logogram_cluster(logo_dict, paragraph) # eg. for ligatures
            elif "x" in logo_dict:
                self._add_ellipsis(logo_dict, paragraph)
            elif "q" in logo_dict:
                # TODO add dedicated function for this
                d_dict = {
                    "frag": logo_dict["q"].replace("|", ""),
                    "delim": logo_dict.get("delim"),
                }
                self._add_excised_d_node(d_dict, paragraph)
                print_if_verbose("Added qualified element {0}".format(logo_dict["q"]))
            elif "c" in logo_dict:
                # TODO: add dedicated function for this
                c_frag = logo_dict["c"].replace("|", "")
                paragraph.add_run(c_frag + logo_dict.get("delim"))
                print_if_verbose("Added composite fragment {0}".format(c_frag))
            elif "mods" in logo_dict:
                # TODO add dedicated function for this
                self._add_pre_frag_symbols(logo_dict, paragraph)
                frag = logo_dict["form"]
                r = paragraph.add_run(frag)
                if frag.islower():
                    r.italic = True
                self._add_post_frag_symbols(logo_dict, paragraph)
                paragraph.add_run(logo_dict.get("delim", ""))
                print_if_verbose("Added MODS logo cluster {0}".format(logo_dict["form"]))
            else:
                print_if_verbose("Non-sign or determinative found in logogram cluster {0}".format(logo_dict))
        paragraph.add_run(gdl_node.get("delim", "")) # delim after the cluster

    def _add_ellipsis(self, gdl_node, paragraph):
        """Adds in things like (...), [...]
        """
        assert(gdl_node.get("x") == "ellipsis")

        self._add_pre_frag_symbols(gdl_node, paragraph)
        paragraph.add_run("...")
        self._add_post_frag_symbols(gdl_node, paragraph)

    def _add_number(self, gdl_node, paragraph):
        """Adds a number to current paragraph, eg. 4.
        """
        self._add_pre_frag_symbols(gdl_node, paragraph)
        num = gdl_node["form"]
        if num == "1/2":
            num = "½"
        elif num == "1/3":
            num = "⅓"
        elif num == "2/3":
            num = "⅔"
        paragraph.add_run(num)
        self._add_post_frag_symbols(gdl_node, paragraph)

        #print_if_verbose("Added number {0}".format(gdl_node["form"]))

    def _add_pre_frag_symbols(self, gdl_node, paragraph):
        """Adds any symbols that come before the actual text fragment.
        These chars may be added: [ ⸢ < <<
        Args:
            gdl_node (dict): dict of L(emma) node's "gdl" property TODO not for det!
            paragraph (docx.text.paragraph.Paragraph): paragraph to add to
        """
        # Full fragment break start
        if gdl_node.get("breakStart", ""):
            paragraph.add_run("[")

        # o for whatever reason may include [ or ], but this is already taken
        # care of by breakStart and breakEnd. Leave o to just be eg. ( ) < >>
        o_frag = gdl_node.get("o", "").strip("[]")
        if "id" in gdl_node and gdl_node.get("statusStart", "") == gdl_node["id"]:
            # o needs to be mirrored first to be an opener frag like ( or < or <<
            o_frag = o_frag.replace("<<", "«").replace("<", "‹").replace(">>", "»").replace(">", "›").replace("$", "")
            o_mirror = closing_punct_mirror[o_frag]
            paragraph.add_run(o_mirror)

        elif gdl_node.get("statusStart", "") == 1:
            # o should already be an opener frag like ( or < or <<
            paragraph.add_run(o_frag)

        # Partial fragment break start
        if gdl_node.get("ho", ""):
            paragraph.add_run("⸢") # note: may look inverted on P50, but it's normal, I assure you

    def _add_post_frag_symbols(self, gdl_node, paragraph):
        """Adds any symbols that come after the actual text fragment.
        These chars may be added: ? ⸣ > >>
        Args:
            gdl_node (dict): dict of L(emma) node's "gdl" property
            paragraph (docx.text.paragraph.Paragraph): paragraph to add to
        """
        # Unknown/uncertain sign
        if gdl_node.get("queried", ""):
            r = paragraph.add_run("(?)")
            #r.font.superscript = True

        # Partial fragment break end
        if gdl_node.get("hc", ""):
            paragraph.add_run("⸣") # note: may look inverted on P50, but it's normal, I assure you

        # o for whatever reason may include [ or ], but this is already taken
        # care of by breakStart and breakEnd. Leave o to just be eg. ( ) < >>
        o_frag = gdl_node.get("o", "").strip("[]")
        if "id" in gdl_node and gdl_node.get("statusStart", "") == gdl_node["id"]:
            # o should already be a closer frag like ) or > or >>
            o_frag = o_frag.replace("<<", "«").replace("<", "‹").replace(">>", "»").replace(">", "›").replace("$", "")
            paragraph.add_run(o_frag)

        # Full fragment break end
        if gdl_node.get("breakEnd", ""):
            paragraph.add_run("]")

        # Whatever delimiter follows, eg. - or space
        if gdl_node.get("delim", ""):
            if gdl_node.get("delim") == "/": # eg. AB / BA needs spacing around / to parse correctly
                paragraph.add_run(" {0} ".format(gdl_node.get("delim")))
            else:
                paragraph.add_run(gdl_node["delim"])

    def _convert_2_or_3_subscript(self, sign):
        """Converts a sign containing a numerical 2 or 3 subscript to have its
        first vowel be properly accented.
        NOTE: Technically not needed...
        Args:
            sign (str): transliterated string of sign with some subscript,
                 eg. bi₂
        Returns:
            str: transliterated string of sign with subscript transformed into
                 accented mark, eg. bí for above example
        """
        if not sign[-1].isdigit(): # No subscript # here
            return sign
        if len(sign) > 2 and sign[-2].isdigit():  # number was eg. 12 or 13- not just 2 or 3
            return sign

        if sign[-1] == "₂":
            subscript_num = "two"
        elif sign[-1] == "₃":
            subscript_num = "three"
        else:
            return sign

        # First vowel in sign will get the accent mark
        for char in sign:
            char_lower = char.lower()
            if char_lower in "aeiu":
                accented_char = vowel_subscript_to_accent_map[char_lower][subscript_num]
                if char.isupper():
                    accented_char = accented_char.upper()
                return sign[:-1].replace(char, accented_char, 1)
        print_if_verbose("You shouldn't be here!")
        return sign

    def _convert_h(self, sign):
        """Replaces h with ḫ, capital or lowercase.
        NOTE You don't need this, since OCHRE can auto sub in rocker-h for regular h
        Args:
            sign (str): transliterated sign, eg. ah or HA
        """
        if sign.islower():
            return sign.replace("h", "ḫ")
        else:
            return sign.replace("H", "Ḫ")


class HtmlParser(object):
    """
    Class to take in a remote link to an ORACC HTML text and output a well-formatted docx.
    Not recommended for RINAP because its superscript Sumerian signs are
    not tagged with <sup> at all.

    TODO: add functionality to take in offline/local HTML file as well?
    """
    def __init__(self, url, catalogue_path):
        self.original_url = url
        self.catalogue_path = catalogue_path

        self.q_number = self._get_q_number()
        self.html_text = requests.get(url).content # TODO give this its own function later for extra checks for right URL
        self.catalogue_dict = self._load_catalogue(catalogue_path)

    def _get_q_number(self):
        """Get the Q-number of the text referred to by the original URL.
        Returns:
            q_number (str)
        """
        url = self.original_url.strip("/")
        q_number = url.split("/")[-1]
        assert("Q" in q_number)
        return q_number

    def _load_catalogue(self):
        """Get catalogue info such as shorthand name and sources list from the JSON
        catalogue passed in for the given Q-number.
        Important keys:
        - display_name
        - exemplars
        - collection (in case exemplars have no museum #)
        - designation/primary_publication (shorthand names in ORACC, put after display_name)
        """
        catalogue_dict = {}

        try:
            with open(self.catalogue_path, encoding="utf_8_sig") as fd:
                raw_str = fd.read()
                catalogue_dict = json.loads(raw_str)
        except Exception as e:
            raise e("Catalogue at {0} not found or unreadable: {1}".format(self.catalogue_path, e))

        try:
            catalogue_dict = catalogue_dict["members"][self.q_number]
        except IndexError as e:
            raise e("Q-number {0} not found in {1}! Did you pick the correct catalogue path?".format(self.q_number, self.catalogue_path))

        return catalogue_dict

    def scrape_page(self):
        """Gets the contents of a given ORACC text URL, eg. http://oracc.museum.upenn.edu/rinap/rinap1/Q003414/html
        and returns it in formatted docx form.
        Returns:
            docx.Document: docx file containing properly formatted text
        """
        doc = Document()
        soup = BeautifulSoup(self.html_text, "html.parser")
        # Preface with name, sources
        # Find all headers
        # Find all text underneath
        soup.find_all("div", class_="")  # TODO
        # ^The most parent node is a tabel with class "transliteration"
        #   1) tr with class eg. "l p2"
        # Drill into a tr with id="" class="h surface";
        #   get span with class="h2 " contents to get eg. Obverse
        # For each tr with class "l" and has id eg. Q003803.1;
        #   go to td with class "lnum ";
        #     and span with class "lnum " or "xlabel " contents are line number, eg. 1
        #   go to td with class "tlit";
        #     every span with class "w *" is a word with a space after it;
        #       get whatever stuff precedes the below 3 (eg. ( 30- ) right below a with class "cbd " directly below the "w " span (this may be the only thing in word-span),
        #       every span with class "r " contains some ... ] thing
        #       every span with class "sign sux " contains a Sumerian word
        #       every span with class "akk " contains an Akkadian sign
        #       every sup with class " " or "sux " (used for ? ! superscript) or "akk " (is then in italics) contains a superscript. This can contain further sups TODO: what about *?
        #       every span with class "compound " contains spans with class "sign " whose contents can be added along with anything before/after them
        #       and get whatever stuff (eg. - . 30- ⸢ -⸣) following any of the above 3
        #   Skip English sections in a td with class eg. "t1 xtr" if any;

    def _add_sumerian_sign(self, html_class, html_text, paragraph):
        """Add a span with class "sign sux ".
        Convert all ng to ^g.
        If there's another sumerian sign right after this, . as separator as default.
        Most come capitalized already, some are not though. Autocorrect or add as-is?
        """
        pass

    def _add_superscript(self, html_class, html_text, paragraph):
        """Add a sup with class " ".
        Convert all ng to ^g.
        If there's another superscript right after this, . separator as default? (If any separator, it could be bundled with actual contents of sup, not after sup)
        """
        pass

    def _add_akkadian_sign(self, html_class, html_text, paragraph):
        """Add a span with class "akk ".
        Comes after an a with class "cbd " that can be ignored.
        Convert all ng to ^g.
        """
        pass

    def _add_pre_span_chars(self, text, paragraph):
        """Add whatever's in between each word-span's a with class "cbd "
        and the first span (one of 3 types) underneath said a. Convert all <,
        >, <<, >> to their Unicode counterparts.
        """
        pass

    def save_docx(self, docx, save_path=None):
        """Save docx to specified location. If none is specified, save to current directory
        under the name of the text's shorthand name, eg. "RINAP 4 029.docx".
        Args:
            docx (docx.Document): transliterated and formatted docx file to be saved
            save_path (str): Path to the directory in which this docx file will be saved.
        """
        if not save_path:
            save_path = os.getcwd()
        if not os.path.isdir(save_path):
            print("Specified save path {0} is not a directory, saving file in current directory instead".format(save_path))
            save_path = os.getcwd()  # TODO not the most obvious... better to try a directory below if existing?

    def get_docx_title(self, docx):
        """Get the name to save the current docx as.
        TODO where would I get it from online without resorting to the ORACC JSONs? The saving functionality should be
        contained with another object, not implemented tbh
        """
        pass


def main():
    """
    Parses arguments, determines which mode (json or html) to use.

    json: [--file /path/to/json] [--directory /path (. by default)]
    html: [--file /path/to/html] [--]
    """
    parser = argparse.ArgumentParser(description='Parse your JSON here.')
    parser.add_argument('--file', '-f', required=True,
                        help='A path (file or directory) to the JSON file to parse into DOCX')
    parser.add_argument('--verbose', '-v', required=False, action="store_true",
                        help='Enable verbose mode during parsing')
    parser.add_argument('--output-directory', '-o', required=False, action="store", default=".",
                        help="Specify directory to output result(s) to. This script will output to the current directory by default.")
    args = parser.parse_args()

    if args.verbose:
        global VERBOSE_FLAG
        VERBOSE_FLAG = True

    jl = JsonLoader(args.file)
    for json_dict in jl.get_json_dicts():
        jp = JsonParser(json_dict, args.output_directory)
        jp.run()


if __name__ == "__main__":
    main()
